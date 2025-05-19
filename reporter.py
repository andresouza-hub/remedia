from docx import Document

def generate_report(desc_df, imagens):
    doc = Document()
    doc.add_heading("Relatório Técnico – RemedIA", level=1)
    doc.add_paragraph("Este relatório apresenta as análises estatísticas e gráficos gerados pelo RemedIA.")

    # Estatísticas Descritivas
    doc.add_heading("Estatísticas Descritivas", level=2)
    doc.add_paragraph(desc_df.to_string())

    # Gráficos
    doc.add_heading("Gráficos", level=2)
    for img_path in imagens:
        doc.add_paragraph("Gráfico gerado:")
        doc.add_picture(img_path, width=5000000)

    output_path = "Relatorio_RemedIA.docx"
    doc.save(output_path)
    return output_path
